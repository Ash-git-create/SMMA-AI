"""Append verified DOI (https://doi.org/...) or canonical URL to each reference,
in both references.md and the report.docx bibliography (in place)."""
import re
from pathlib import Path
from docx import Document

ROOT = Path(r"D:/Master Thesis/SMMA_AI_Systems")
REFS = ROOT / "docs" / "writing" / "references.md"
DOC = ROOT / "report.docx"

LINKS = {
    1: "https://arxiv.org/abs/2308.08155",
    2: "https://arxiv.org/abs/2308.00352",
    3: "https://arxiv.org/abs/2404.16130",
    4: "https://doi.org/10.1145/3571730",
    5: "https://arxiv.org/abs/2607.21912",
    6: "https://arxiv.org/abs/2603.04474",
    7: "https://arxiv.org/abs/2402.01680",
    8: "https://proceedings.neurips.cc/paper/2020/hash/6b493230205f780e1bc26945df7481e5-Abstract.html",
    9: "https://doi.org/10.1145/3703155",
    10: "https://doi.org/10.18653/v1/2020.acl-main.173",
    11: "https://doi.org/10.1098/rspa.1927.0118",
    12: "https://doi.org/10.1137/S0036144500371907",
    13: "https://doi.org/10.1561/1900000006",
    14: "http://www.cidrdb.org/cidr2005/papers/P22.pdf",
    15: "http://www.vldb.org/conf/2006/p953-benjelloun.pdf",
    16: "https://arxiv.org/abs/2402.07867",
    17: "https://arxiv.org/abs/2606.07937",
    18: "https://arxiv.org/abs/2606.20493",
    19: "https://arxiv.org/abs/2606.24535",
    20: "https://arxiv.org/abs/2606.27409",
    21: "https://doi.org/10.18653/v1/2020.emnlp-main.550",
    22: "https://doi.org/10.1109/TKDE.2024.3352100",
    23: "https://doi.org/10.18653/v1/2023.emnlp-main.557",
    24: "https://proceedings.neurips.cc/paper_files/paper/2023/hash/91f18a1287b398d378ef22505bf41832-Abstract-Datasets_and_Benchmarks.html",
    25: "https://doi.org/10.1103/PhysRevLett.86.3200",
    26: "https://doi.org/10.1103/PhysRevE.66.016128",
    27: "https://doi.org/10.1109/ICDE.2008.4497511",
}

# --- references.md ---
content = REFS.read_text(encoding="utf-8")
blocks = content.split("\n\n")
out = []
for b in blocks:
    m = re.match(r"^\[(\d+)\]", b.strip())
    if m and int(m.group(1)) in LINKS and "http" not in b:
        b = b.rstrip() + " Available: " + LINKS[int(m.group(1))]
    out.append(b)
REFS.write_text("\n\n".join(out), encoding="utf-8")

# --- report.docx bibliography (only paragraphs after the Bibliography heading) ---
d = Document(str(DOC))
in_bib = False
added = 0
for p in d.paragraphs:
    if p.style.name.startswith("Heading") and p.text.strip() == "Bibliography":
        in_bib = True
        continue
    if p.style.name.startswith("Heading") and in_bib:
        in_bib = False  # left the bibliography (e.g., List of Figures)
    if not in_bib:
        continue
    m = re.match(r"^\[(\d+)\]", p.text.strip())
    if m and int(m.group(1)) in LINKS and "http" not in p.text:
        link = "  Available: " + LINKS[int(m.group(1))]
        if p.runs:
            p.runs[-1].text = p.runs[-1].text.rstrip() + link
        else:
            p.add_run(link)
        added += 1
d.save(str(DOC))
print("references.md updated; docx bibliography links added:", added)
