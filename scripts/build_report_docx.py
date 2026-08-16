"""Build report.docx from the SRH template + approved chapter markdown.

Inherits the template's styles/margins/headers, clears the placeholder body, and
writes: title page, bilingual affidavit, English abstract + German Zusammenfassung,
an auto-updating Table of Contents field, the approved Chapter 1 sections, and the
references cited so far. Re-run after each newly approved section.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(r"D:/Master Thesis/SMMA_AI_Systems")
TEMPLATE = ROOT / "For writing" / "Template_Bachelor_Master_EN.docx"
CH1 = ROOT / "docs" / "chapters" / "ch1_introduction.md"
CH2 = ROOT / "docs" / "chapters" / "ch2_state_of_the_art.md"
CH3 = ROOT / "docs" / "chapters" / "ch3_methodology.md"
REFS = ROOT / "docs" / "writing" / "references.md"
OUT = ROOT / "report.docx"

TITLE = "Cascading Knowledge Contamination in Shared Memory Multi-Agent AI Systems"
AUTHOR = "Ashwin Jayan"
MATRIC = "100002367"
DATE = "August 2026"
UNIVERSITY = "SRH University Heidelberg"
SCHOOL = "School of Information, Media and Design"
COURSE = "M.Sc. Applied Data Science and Analytics"
REVIEWER1 = "Prof. Dr. Ing. Binh Vu"
REVIEWER2 = "Prof. Dr. Ing. Swati Chandna"

ABSTRACT_EN = (
    "Multi-agent systems built from large language models increasingly share a common "
    "memory, often a knowledge graph, so that one agent can reuse what another has found. "
    "This thesis studies a failure mode of that design: when an agent writes a wrong fact "
    "into the shared memory, other agents retrieve it as if it were true, build on it, and "
    "write further wrong facts, so the error spreads. The thesis calls this cascading "
    "knowledge contamination and studies it in a controlled testbed. A knowledge graph in "
    "Neo4j is filled with correct facts from the T-REx dataset, agents extract facts from "
    "the HotpotQA and FEVER datasets, and three types of extraction error (entity "
    "disambiguation errors, qualifier loss, and relation strengthening) are injected on "
    "purpose, one at a time. The spread of contamination is tracked with a discrete-time "
    "SIR model borrowed from epidemiology, which yields a basic reproduction number R₀ for "
    "each error type and system setting. The thesis then tests a provenance-aware "
    "mitigation, adapted from the Stanford Trio database system, that records the origin of "
    "each fact and can lower the trust in facts derived from a fact found to be wrong. The "
    "main findings are that contamination spreads only when the error reaches the part of "
    "the graph that agents retrieve, and this reachability behaves as a threshold rather "
    "than a gradient; that the amount of contaminated memory and the harm to answers "
    "decouple, so memory can be badly corrupted while task scores stay flat; that entity "
    "errors spread the most and relation strengthening the least; and that the "
    "provenance-aware mitigation does not contain contamination under a realistic "
    "language-model validator, because the injected errors erase their own contradicting "
    "evidence. The limiting factor is validator recall, not precision, and the failure "
    "reproduces across model family and capability. The thesis contributes a fact-level "
    "epidemiological framework, a set of measured results on how contamination behaves, a "
    "negative result on provenance-aware mitigation together with its mechanism, and an "
    "open testbed for further study."
)

ABSTRACT_DE = (
    "Multi-Agenten-Systeme aus großen Sprachmodellen nutzen zunehmend einen gemeinsamen "
    "Speicher, häufig einen Wissensgraphen, damit ein Agent wiederverwenden kann, was ein "
    "anderer gefunden hat. Diese Arbeit untersucht einen Versagensmodus dieses Entwurfs: "
    "Schreibt ein Agent eine falsche Tatsache in den gemeinsamen Speicher, rufen andere "
    "Agenten sie ab, als wäre sie wahr, bauen darauf auf und schreiben weitere falsche "
    "Tatsachen, sodass sich der Fehler ausbreitet. Die Arbeit nennt dies kaskadierende "
    "Wissenskontamination und untersucht sie in einer kontrollierten Testumgebung. Ein "
    "Wissensgraph in Neo4j wird mit korrekten Tatsachen aus dem T-REx-Datensatz gefüllt, "
    "Agenten extrahieren Tatsachen aus den Datensätzen HotpotQA und FEVER, und drei Arten "
    "von Extraktionsfehlern (Entitäts-Verwechslung, Verlust von Qualifikatoren und "
    "Relationsverstärkung) werden gezielt und einzeln eingefügt. Die Ausbreitung der "
    "Kontamination wird mit einem zeitdiskreten SIR-Modell aus der Epidemiologie verfolgt, "
    "das eine Basisreproduktionszahl R₀ für jeden Fehlertyp und jede Systemkonfiguration "
    "liefert. Anschließend testet die Arbeit eine herkunftsbewusste Gegenmaßnahme, "
    "angelehnt an das Trio-Datenbanksystem der Stanford University, die den Ursprung jeder "
    "Tatsache aufzeichnet und das Vertrauen in davon abgeleitete Tatsachen senken kann, sobald "
    "eine Tatsache als falsch erkannt wird. Die wichtigsten Ergebnisse: Kontamination "
    "breitet sich nur aus, wenn der Fehler den Teil des Graphen erreicht, den Agenten "
    "abrufen, und diese Erreichbarkeit verhält sich wie eine Schwelle und nicht wie ein "
    "Gradient; die Menge des kontaminierten Speichers und der Schaden an den Antworten "
    "entkoppeln sich, sodass der Speicher stark beschädigt sein kann, während die "
    "Antwortgenauigkeit unverändert bleibt; Entitätsfehler breiten sich am stärksten "
    "aus, Relationsverstärkung am wenigsten; und die herkunftsbewusste Gegenmaßnahme "
    "dämmt die Kontamination unter einem realistischen Sprachmodell-Prüfer nicht ein, "
    "weil die eingefügten Fehler ihre eigenen widersprechenden Belege auslöschen. Der "
    "begrenzende Faktor ist die Trefferquote (Recall) des Prüfers, nicht seine "
    "Genauigkeit (Precision), und das Versagen tritt über Modellfamilien und Modellstärken "
    "hinweg auf. Die Arbeit liefert einen epidemiologischen Rahmen auf Faktenebene, eine "
    "Reihe gemessener Ergebnisse zum Verhalten der Kontamination, ein negatives Ergebnis "
    "zur herkunftsbewussten Gegenmaßnahme samt Mechanismus sowie eine offene "
    "Testumgebung für weitere Forschung."
)

AFFIDAVIT_EN = (
    "Herewith I declare: that I have composed the chapters for the Master Thesis for which "
    "I am named as the author independently, that I did not use any other sources or "
    "additives than the ones specified, and that I did not submit this work at any other "
    "examination procedure."
)
AFFIDAVIT_DE = (
    "Ich versichere, dass ich die Kapitel der Arbeit, für die ich als Verfasser genannt "
    "werde, selbständig verfasst habe, dass ich keine anderen als die angegebenen Quellen "
    "und Hilfsmittel benutzt habe und dass ich diese Arbeit bei keinem anderen "
    "Prüfungsverfahren vorgelegt habe."
)

REFERENCES = [
    '[1] Q. Wu, G. Bansal, J. Zhang, Y. Wu, B. Li, et al., "AutoGen: Enabling Next-Gen LLM '
    'Applications via Multi-Agent Conversation," arXiv:2308.08155, 2023.',
    '[2] S. Hong, M. Zhuge, J. Chen, X. Zheng, Y. Cheng, C. Zhang, et al., "MetaGPT: Meta '
    'Programming for a Multi-Agent Collaborative Framework," in Proc. ICLR, 2024; '
    'arXiv:2308.00352.',
    '[3] D. Edge, H. Trinh, N. Cheng, J. Bradley, A. Chao, A. Mody, S. Truitt, et al., '
    '"From Local to Global: A Graph RAG Approach to Query-Focused Summarization," '
    'arXiv:2404.16130, 2024.',
    '[4] Z. Ji, N. Lee, R. Frieske, T. Yu, D. Su, Y. Xu, E. Ishii, Y. J. Bang, A. Madotto, '
    'and P. Fung, "Survey of Hallucination in Natural Language Generation," ACM Computing '
    'Surveys, vol. 55, no. 12, Article 248, pp. 1-38, 2023.',
    '[5] R. Niu, X. Shu, and Y. Zhao, "Reliability-Contagion Feasibility in LLM '
    'Multi-Agent Networks," arXiv:2607.21912, 2026.',
]


ACKNOWLEDGEMENTS = (
    "I thank my supervisors, Prof. Dr. Ing. Binh Vu and Prof. Dr. Ing. Swati Chandna, "
    "for their guidance throughout this thesis. I also thank my family and friends for "
    "their support during the work."
)

ACRONYMS = [
    ("AI", "Artificial Intelligence"),
    ("API", "Application Programming Interface"),
    ("AUROC", "Area Under the Receiver Operating Characteristic curve"),
    ("DNF", "Disjunctive Normal Form"),
    ("EM", "Exact Match"),
    ("FEVER", "Fact Extraction and Verification (dataset)"),
    ("KG", "Knowledge Graph"),
    ("LLM", "Large Language Model"),
    ("MAS", "Multi-Agent System"),
    ("QA", "Question Answering"),
    ("R₀", "Basic Reproduction Number"),
    ("RAG", "Retrieval-Augmented Generation"),
    ("SIR", "Susceptible-Infected-Recovered model"),
    ("SPO", "Subject-Predicate-Object"),
    ("T-REx", "Dataset of Wikipedia text aligned with Wikidata triples"),
    ("ULDB", "Uncertainty-Lineage Database"),
    ("USR", "Unsupported Sentence Ratio"),
]

FIGURES = [
    "Figure 2.1  The SIR model mapped onto knowledge-graph facts",
    "Figure 2.2  Two pictures of contagion: message chain versus shared graph",
    "Figure 3.1  The system architecture",
    "Figure 3.2  The three error injections, shown as before/after triplets",
    "Figure 3.3  The experiment pipeline",
]

TABLES = [
    "Table 2.1  How this thesis compares with the closest recent work",
    "Table 3.1  The three agents, the model each runs, and its role",
    "Table 3.2  The fields stored with every triplet",
    "Table 3.3  The three error types, with an illustrative example of each",
    "Table 3.4  The experiment arms",
    "Table 3.5  The evaluation metrics",
]


def add_kv_table(doc, pairs, headers):
    table = doc.add_table(rows=1, cols=2)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for i, h in enumerate(headers):
        r = table.rows[0].cells[i].paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9)
    for k, v in pairs:
        cells = table.add_row().cells
        rk = cells[0].paragraphs[0].add_run(k)
        rk.bold = True
        rk.font.size = Pt(9)
        rv = cells[1].paragraphs[0].add_run(v)
        rv.font.size = Pt(9)


def clear_body(doc):
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is sectPr:
            continue  # keep section props in the tree so add_table() can size columns
        body.remove(child)
    return body, sectPr


def centered(doc, text, size=None, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    if size:
        r.font.size = Pt(size)
    return p


def normal(doc, text):
    return doc.add_paragraph(text, style="Normal")


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = 'Right-click here and choose "Update Field" to build the contents.'
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)


def load_references():
    text = REFS.read_text(encoding="utf-8")
    refs = []
    for block in re.split(r"\n\s*\n", text):
        b = block.strip()
        if re.match(r"^\[\d+\]", b):
            refs.append(" ".join(ln.strip() for ln in b.split("\n")))
    return refs


def add_chapter_md(doc, path):
    text = path.read_text(encoding="utf-8")
    # drop HTML comments
    text = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)
    blocks = re.split(r"\n\s*\n", text)
    for block in blocks:
        lines = [ln for ln in block.split("\n") if ln.strip() != ""]
        if not lines:
            continue
        first = lines[0].strip()
        fig = re.match(r"^!\[(?P<cap>.*)\]\((?P<path>.*)\)$", first)
        if fig and len(lines) == 1:
            pic = doc.add_paragraph()
            pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic.add_run().add_picture(str(ROOT / fig.group("path")), width=Inches(5.8))
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cap.add_run(fig.group("cap"))
            cr.italic = True
            cr.font.size = Pt(9)
            continue
        if first.startswith("[TBL]"):
            cap = doc.add_paragraph()
            cr = cap.add_run(first[5:].strip())
            cr.italic = True
            cr.font.size = Pt(9)
            continue
        if first.startswith("|") and len(lines) >= 2 and set(lines[1].strip()) <= set("|-: "):
            rows = [ln for ln in lines if ln.strip().startswith("|")]
            header = [c.strip() for c in rows[0].strip().strip("|").split("|")]
            table = doc.add_table(rows=1, cols=len(header))
            try:
                table.style = "Table Grid"
            except Exception:
                pass
            for i, h in enumerate(header):
                run = table.rows[0].cells[i].paragraphs[0].add_run(h)
                run.bold = True
                run.font.size = Pt(9)
            for r in rows[2:]:
                cells = [c.strip() for c in r.strip().strip("|").split("|")]
                rowcells = table.add_row().cells
                for i in range(len(header)):
                    run = rowcells[i].paragraphs[0].add_run(cells[i] if i < len(cells) else "")
                    run.font.size = Pt(9)
            continue
        if first.startswith("# "):
            continue  # chapter title added manually
        if first.startswith("### "):
            doc.add_heading(first[4:].strip(), level=3)
            continue
        if first.startswith("## "):
            doc.add_heading(first[3:].strip(), level=2)
            continue
        joined = " ".join(ln.strip() for ln in lines)
        if joined.startswith("**"):
            close = joined.find("**", 2)
            bold_text = joined[2:close].strip()
            rest = joined[close + 2:].strip()
            p = doc.add_paragraph(style="Normal")
            r = p.add_run(bold_text)
            r.bold = True
            if rest:
                doc.add_paragraph(rest, style="Normal")
        else:
            doc.add_paragraph(joined, style="Normal")


def main():
    doc = Document(str(TEMPLATE))
    body, sectPr = clear_body(doc)

    # --- Title page ---
    for _ in range(3):
        doc.add_paragraph()
    centered(doc, TITLE, size=22, bold=True, space_after=24)
    centered(doc, "Master Thesis", size=16)
    centered(doc, "by")
    centered(doc, AUTHOR, size=14, bold=True)
    centered(doc, f"Matriculation number: {MATRIC}", space_after=24)
    centered(doc, DATE, space_after=24)
    centered(doc, UNIVERSITY, size=13)
    centered(doc, SCHOOL)
    centered(doc, COURSE, space_after=24)
    centered(doc, "Reviewers")
    centered(doc, REVIEWER1)
    centered(doc, REVIEWER2)
    doc.add_page_break()

    # --- Affidavit ---
    doc.add_heading("Affidavit", level=1)
    normal(doc, AFFIDAVIT_EN)
    normal(doc, "Heidelberg, ____________ (date)      Signature ______________________")
    doc.add_paragraph()
    doc.add_heading("Ehrenwörtliche Erklärung", level=1)
    normal(doc, AFFIDAVIT_DE)
    normal(doc, "Heidelberg, ____________ (Datum)     Unterschrift ___________________")
    doc.add_page_break()

    # --- Acknowledgement ---
    doc.add_heading("Acknowledgement", level=1)
    normal(doc, ACKNOWLEDGEMENTS)
    doc.add_page_break()

    # --- Abstract ---
    doc.add_heading("Abstract", level=1)
    normal(doc, ABSTRACT_EN)
    doc.add_heading("Zusammenfassung", level=1)
    normal(doc, ABSTRACT_DE)
    doc.add_page_break()

    # --- Table of Contents ---
    doc.add_heading("Table of Contents", level=1)
    add_toc(doc)
    doc.add_page_break()

    # --- List of Acronyms ---
    doc.add_heading("List of Acronyms", level=1)
    add_kv_table(doc, ACRONYMS, ("Abbreviation", "Meaning"))
    doc.add_page_break()

    # --- Chapter 1 ---
    doc.add_heading("1  Introduction", level=1)
    add_chapter_md(doc, CH1)

    # --- Chapter 2 ---
    doc.add_page_break()
    doc.add_heading("2  State of the Art", level=1)
    add_chapter_md(doc, CH2)

    # --- Chapter 3 ---
    doc.add_page_break()
    doc.add_heading("3  Methodology", level=1)
    add_chapter_md(doc, CH3)

    # --- References (cited so far) ---
    doc.add_page_break()
    doc.add_heading("Bibliography", level=1)
    for ref in load_references():
        doc.add_paragraph(ref, style="Normal")

    # --- List of Figures ---
    doc.add_page_break()
    doc.add_heading("List of Figures", level=1)
    for item in FIGURES:
        doc.add_paragraph(item, style="Normal")

    # --- List of Tables ---
    doc.add_page_break()
    doc.add_heading("List of Tables", level=1)
    for item in TABLES:
        doc.add_paragraph(item, style="Normal")

    if sectPr is not None:
        body.remove(sectPr)
        body.append(sectPr)  # ensure section props are the last body child

    # Tell Word to refresh all fields (Table of Contents page numbers) on open.
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    doc.settings.element.insert(0, upd)

    doc.save(str(OUT))
    print(f"Wrote {OUT}")
    print(f"Paragraphs: {len(doc.paragraphs)}")


if __name__ == "__main__":
    main()
