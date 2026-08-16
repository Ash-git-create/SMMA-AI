"""In-place: replace affidavit with the SRH Declaration of Authorship, expand the
Acknowledgements, and add an honest AI-use declaration. Preserves manual formatting.
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

DOC = r"D:/Master Thesis/SMMA_AI_Systems/report.docx"

DECL1 = ("I hereby declare that my herewith submitted thesis is my own original work. I "
         "have written it independently without outside help and have not used any sources "
         "other than those indicated - in particular, no sources not named in the "
         "references.")
DECL2 = ("I have appropriately indicated any direct quotations or passages taken from "
         "literature, and the use of intellectual property from other authors, by "
         "providing the necessary citations within the work. This applies equally to the "
         "sources used for text generation by Artificial Intelligence (AI).")
DECL3 = ("I hereby declare that the thesis was not previously presented to another "
         "examination board, and I also confirm that the PDF version of this thesis is "
         "exactly identical in content to the hard copy.")
SIG1 = "____________________________, (place)      ____________ (date)"
SIG2 = "____________________________ (signature)"

AI_TEXT = (
    "In preparing this thesis I used the AI assistant Claude (Anthropic), accessed through "
    "the Claude Code tool, as a writing and programming aid. It was used to help draft and "
    "revise the wording of the text, to produce figures and diagrams from my own "
    "specifications, and to help write, debug, and document parts of the source code. The "
    "design of the research, the experiments and the results they produced, and the "
    "analysis and interpretation of those results are my own. I reviewed all AI-assisted "
    "text and output, checked every fact, number, and citation against the archived "
    "experimental results and the cited sources, and I take full responsibility for the "
    "content of this thesis.")

ACK1 = ("First of all, I would like to express my deep gratitude to my supervisors, Prof. "
        "Dr. Ing. Binh Vu and Prof. Dr. Ing. Swati Chandna, who guided and supported me "
        "throughout the writing of this thesis. Their dedication, expertise, and thoughtful "
        "feedback helped me to shape and complete this work.")
ACK2 = ("I would also like to thank the teachers of the Applied Data Science and Analytics "
        "programme at SRH University Heidelberg, who gave me the foundations and the deeper "
        "knowledge in this field that made this thesis possible.")
ACK3 = ("I am grateful to my fellow students and friends for the many discussions that "
        "sharpened my thinking, and for their encouragement during the more difficult "
        "stretches of the work.")
ACK4 = ("Finally, I would like to thank my family for their constant support and patience "
        "throughout my studies and this research. Their encouragement gave me the "
        "conditions to focus on this thesis and see it through.")

d = Document(DOC)


def set_text(p, text):
    runs = list(p.runs)
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def insert_after(ref, text="", style=None):
    new_p = OxmlElement("w:p")
    ref._p.addnext(new_p)
    np = Paragraph(new_p, ref._parent)
    if style:
        np.style = d.styles[style]
    if text:
        np.add_run(text)
    return np


def find(pred):
    for p in d.paragraphs:
        if pred(p.text.strip()):
            return p
    return None


h_aff = find(lambda t: t == "Affidavit")
p_aff_en = find(lambda t: t.startswith("Herewith I declare"))
p_sig_en = find(lambda t: t.startswith("Heidelberg,") and "(date)" in t)
h_ehren = find(lambda t: t.startswith("Ehrenw"))
p_aff_de = find(lambda t: t.startswith("Ich versichere"))
p_sig_de = find(lambda t: t.startswith("Heidelberg,") and "(Datum)" in t)
h_ack = find(lambda t: t == "Acknowledgement")
p_ack = find(lambda t: t.startswith("I thank my supervisors"))

# --- Declaration of Authorship ---
set_text(h_aff, "Declaration of Authorship")
set_text(p_aff_en, DECL1)
set_text(p_sig_en, DECL2)
insert_after(p_sig_en, DECL3, "Normal")
# repurpose the German heading + paragraphs as the signature block
h_ehren.style = d.styles["Normal"]
set_text(h_ehren, SIG1)
for r in h_ehren.runs:
    r.bold = None
    r.font.size = None
set_text(p_aff_de, SIG2)
set_text(p_sig_de, "")

# --- AI-use declaration (own page) ---
ai_h = insert_after(p_sig_de, "Declaration on the Use of Artificial Intelligence", "Heading 1")
ai_h.paragraph_format.page_break_before = True
insert_after(ai_h, AI_TEXT, "Normal")

# --- Acknowledgements ---
set_text(h_ack, "Acknowledgements")
set_text(p_ack, ACK1)
a = insert_after(p_ack, ACK2, "Normal")
a = insert_after(a, ACK3, "Normal")
insert_after(a, ACK4, "Normal")

d.save(DOC)
print("declaration replaced, AI declaration added, acknowledgements expanded")
