"""In-place surgical fixes to report.docx (preserving manual formatting).

- Remove automatic heading numbering (front/back matter should be unnumbered;
  chapters keep their typed numbers 1, 1.1, 2.1.1, ...).
- Apply the five German fidelity fixes to the Zusammenfassung.
- Fix the Figure 3.3 caption (audit is conditional).
- Swap the pipeline figure image for the corrected one, keeping aspect ratio.
- Ensure Word refreshes the TOC on open.
"""
import struct
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(r"D:/Master Thesis/SMMA_AI_Systems")
DOC = ROOT / "report.docx"

GERMAN_FIXES = [
    ("eine Schwachstelle dieses Entwurfs", "einen Versagensmodus dieses Entwurfs"),
    ("in abgeleitete Tatsachen senken", "in davon abgeleitete Tatsachen senken"),
    ("und nicht wie ein Verlauf", "und nicht wie ein Gradient"),
    ("verhindert die Kontamination unter einem realistischen Sprachmodell-Prüfer nicht",
     "dämmt die Kontamination unter einem realistischen Sprachmodell-Prüfer nicht ein"),
    ("Modellfamilien und -größen", "Modellfamilien und Modellstärken"),
]
CAPTION_FIX = (
    "write facts back, and then measures the outcome",
    "write facts back (and, in arms with validation enabled, audit and quarantine), "
    "and then measures the outcome",
)

doc = Document(str(DOC))

# 1. strip numbering from heading STYLES
for sname in ("Heading 1", "Heading 2", "Heading 3"):
    el = doc.styles[sname].element
    pPr = el.find(qn("w:pPr"))
    if pPr is not None:
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            pPr.remove(numPr)

# 2. strip direct numbering from heading PARAGRAPHS
n_para_num = 0
for p in doc.paragraphs:
    if not p.style.name.startswith("Heading"):
        continue
    pPr = p._p.pPr
    if pPr is not None:
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            pPr.remove(numPr)
            n_para_num += 1

# 3. German fixes (single-run Zusammenfassung paragraph)
ger = [p for p in doc.paragraphs if "Wissenskontamination" in p.text]
for p in ger:
    for r in p.runs:
        for a, b in GERMAN_FIXES:
            if a in r.text:
                r.text = r.text.replace(a, b)

# 4. caption fix
for p in doc.paragraphs:
    for r in p.runs:
        if CAPTION_FIX[0] in r.text:
            r.text = r.text.replace(*CAPTION_FIX)

# 5. swap pipeline image (identify via its caption's preceding paragraph)
paras = doc.paragraphs
cap_idx = next((i for i, p in enumerate(paras) if "The experiment pipeline" in p.text), None)
img_swapped = False
if cap_idx is not None:
    img_para = paras[cap_idx - 1]
    blips = img_para._p.findall(".//" + qn("a:blip"))
    if blips:
        rId = blips[0].get(qn("r:embed"))
        new = (ROOT / "docs" / "figures" / "fig_pipeline.png").read_bytes()
        w, h = struct.unpack(">II", new[16:24])
        for shp in doc.inline_shapes:
            blip = shp._inline.graphic.graphicData.pic.blipFill.blip
            if blip.get(qn("r:embed")) == rId:
                cx = shp.width
                shp.height = int(cx * h / w)
                break
        doc.part.related_parts[rId]._blob = new
        img_swapped = True

# 6. ensure fields refresh on open
sett = doc.settings.element
if sett.find(qn("w:updateFields")) is None:
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    sett.insert(0, upd)

doc.save(str(DOC))
print(f"heading-paragraph numbering stripped: {n_para_num}")
print(f"german paragraphs touched: {len(ger)}")
print(f"pipeline image swapped: {img_swapped}")
