"""Replace spelled-out SIR symbols with proper Unicode: beta->β, gamma->γ, R0->R₀.
Whole-word only. Applies to report.docx (paragraphs + tables), the chapter sources,
and the build-script constants."""
import re
from pathlib import Path
from docx import Document

ROOT = Path(r"D:/Master Thesis/SMMA_AI_Systems")


def sub_all(text):
    text = re.sub(r"\bbeta\b", "\u03b2", text)
    text = re.sub(r"\bgamma\b", "\u03b3", text)
    text = re.sub(r"\bR0\b", "R\u2080", text)
    return text


for rel in ("docs/chapters/ch2_state_of_the_art.md",
            "docs/chapters/ch3_methodology.md",
            "scripts/build_report_docx.py"):
    fp = ROOT / rel
    fp.write_text(sub_all(fp.read_text(encoding="utf-8")), encoding="utf-8")

doc = Document(str(ROOT / "report.docx"))


def fix_runs(paras):
    n = 0
    for p in paras:
        for r in p.runs:
            new = sub_all(r.text)
            if new != r.text:
                r.text = new
                n += 1
    return n


changed = fix_runs(doc.paragraphs)
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            changed += fix_runs(cell.paragraphs)
doc.save(str(ROOT / "report.docx"))
print("docx runs changed:", changed)
