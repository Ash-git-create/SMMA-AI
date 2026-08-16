# -*- coding: utf-8 -*-
"""Insert one paragraph into 3.8 disclosing the calibration's single rater and the
per-category (stratified) correction. Grounded in phase34_judge_calibration_summary.json
(single author rater; 56*0.0 + 32*0.2 + 1 ~= 7.4/783 ~= 0.9%)."""
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from pathlib import Path

DOC = Path(r"D:/Master Thesis/SMMA_AI_Systems/report.docx")
d = Document(str(DOC))

def insert_after(ref, text, style="Normal"):
    p = OxmlElement("w:p"); ref._p.addnext(p)
    np = Paragraph(p, ref._parent); np.style = d.styles[style]; np.add_run(text)
    return np

NOTE = ("Two features of this calibration bound how far it can be pushed. The forty items "
        "were labelled by the author, blind to the judge's verdicts but by a single rater, so "
        "no second-rater agreement can be reported and the corrected figure is best read as a "
        "careful single-annotator estimate with a wide margin. The sample was also stratified "
        "rather than random, chosen to cover each of the judge's verdicts and every relation "
        "flag, so the correction is applied to each category on its own: the share of each "
        "flag type that a human confirmed as a real error is applied back to that type's count "
        "in the full audit, and a separate check of the facts the judge had accepted bounds "
        "how many real errors it missed. A second independent rater, and a stated confidence "
        "interval on the corrected rate, would firm this step up further and are noted here as "
        "the natural next step.")

hits = [p for p in d.paragraphs if p.text.strip().endswith("Chapter 5 gives the size of the gap.")]
assert len(hits) == 1, "anchor count = %d" % len(hits)
insert_after(hits[0], NOTE)
d.save(str(DOC))
print("3.8 calibration note inserted")
