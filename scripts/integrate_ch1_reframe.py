# -*- coding: utf-8 -*-
"""Reframe the 1.2 opener: keep a tightened 0.95^10 hook, then pivot to the
shared-memory reinforcement + reachability-threshold mechanism the thesis
actually studies. Replaces one single-run paragraph and adds one after it."""
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from pathlib import Path

DOC = Path(r"D:/Master Thesis/SMMA_AI_Systems/report.docx")
d = Document(str(DOC))

def insert_after(ref, text, style="Normal"):
    p = OxmlElement("w:p"); ref._p.addnext(p)
    np = Paragraph(p, ref._parent); np.style = d.styles[style]; np.add_run(text)
    return np

P1 = ("The reason a cascade is worth worrying about is that errors in a chain compound "
      "rather than add up. A quick illustration makes the point. Suppose every step in a "
      "ten-step pipeline is right 95% of the time, and each step builds on the one before it. "
      "Then the whole chain is right only about 60% of the time, since 0.95 to the tenth "
      "power is 0.599. A per-step error rate that looks tiny on its own stops being tiny once "
      "the steps depend on each other.")
P2 = ("A shared memory makes this worse in a way a simple chain does not. In a chain each "
      "output is passed on once and then left behind, so an early error tends to be diluted "
      "by fresh input at each later step. In a shared knowledge graph the same contaminated "
      "fact can be retrieved again and again, by many agents and at many later steps, so it "
      "is reused and reinforced instead of fading out. Whether an error spreads at all turns "
      "on one thing: whether it lands in the part of the memory that agents actually read "
      "from. An error that sits in a corner no agent retrieves goes nowhere, however wrong it "
      "is, while an error in a well-read region can be copied into fact after fact. This "
      "thesis measures where that line falls, and how fast contamination moves once it is "
      "crossed.")

hits = [p for p in d.paragraphs if p.text.strip().startswith("The reason a cascade is worth worrying about")]
assert len(hits) == 1, "anchor count = %d" % len(hits)
p = hits[0]
assert len(p.runs) == 1, "expected single run, got %d" % len(p.runs)
p.runs[0].text = P1
insert_after(p, P2)
d.save(str(DOC))
print("1.2 reframing integrated")
