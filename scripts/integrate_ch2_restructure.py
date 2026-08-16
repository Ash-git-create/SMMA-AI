# -*- coding: utf-8 -*-
"""Rebuild 2.2 into three labelled subsections (2.2.1/2.2.2/2.2.3) and add the
convergent-evidence paragraph to 2.3.3. Adds refs [42]-[44]. Gated (style +
fact + originality). In-text citations use clickable bookmark-anchor links
(targets ref5,6,16,17,18,19,20 already exist; ref42-44 created here). The four
old 2.2 body paragraphs are removed after the new ones are inserted."""
import re
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

DOC = Path(r"D:/Master Thesis/SMMA_AI_Systems/report.docx")
d = Document(str(DOC)); part = d.part

def run_like(base_r, text, color=None, underline=False):
    r = OxmlElement("w:r")
    if base_r is not None and base_r.find(qn("w:rPr")) is not None:
        r.append(deepcopy(base_r.find(qn("w:rPr"))))
    if color or underline:
        rpr = r.find(qn("w:rPr"))
        if rpr is None: rpr = OxmlElement("w:rPr"); r.insert(0, rpr)
        if color: c = OxmlElement("w:color"); c.set(qn("w:val"), color); rpr.append(c)
        if underline: u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    r.append(t); return r

def hyperlink(child, anchor=None, rid=None):
    hl = OxmlElement("w:hyperlink")
    if anchor: hl.set(qn("w:anchor"), anchor)
    if rid: hl.set(qn("r:id"), rid)
    hl.append(child); return hl

def insert_after(ref, text, style="Normal"):
    p = OxmlElement("w:p"); ref._p.addnext(p)
    np = Paragraph(p, ref._parent); np.style = d.styles[style]; np.add_run(text)
    return np

def linkify_run(run):
    base = run._r
    if not re.search(r"\[\d+\]", run.text): return
    new = []
    for pc in re.split(r"(\[\d+\])", run.text):
        m = re.fullmatch(r"\[(\d+)\]", pc)
        if m: new.append(hyperlink(run_like(base, pc), anchor="ref" + m.group(1)))
        elif pc: new.append(run_like(base, pc))
    parent = base.getparent(); idx = list(parent).index(base); parent.remove(base)
    for i, el in enumerate(new): parent.insert(idx + i, el)

def link_intext(paragraph):
    for run in list(paragraph.runs):
        if re.search(r"\[\d+\]", run.text): linkify_run(run)

def para(pred):
    hits = [p for p in d.paragraphs if pred(p.text.strip())]
    assert len(hits) == 1, "expected 1, got %d for %s" % (len(hits), pred)
    return hits[0]

# ---- locate anchors first ----
heading = para(lambda t: t == "2.2 Related Work" and True) if any(
    p.text.strip()=="2.2 Related Work" and p.style.name=="Heading 2" for p in d.paragraphs) else None
heading = next(p for p in d.paragraphs if p.text.strip()=="2.2 Related Work" and p.style.name=="Heading 2")
old_a = para(lambda t: t.startswith("Three strands of recent work"))
old_b = para(lambda t: t.startswith("The second strand models"))
old_c = para(lambda t: t.startswith("The third strand tries"))
old_d = para(lambda t: t.startswith("These three strands give"))
opt_anchor = para(lambda t: t.startswith("One containment result points the same way"))
bib41 = para(lambda t: t.startswith("[41]"))
print("anchors OK")

INTRO = ("Three strands of recent work touch this thesis: attacks on agent memory, models of "
 "how errors spread between agents, and attempts to contain that spread. This section takes "
 "each in turn, and Section 2.3 sets out where this thesis sits against it.")
B221 = ("The first strand is adversarial. Recent work shows that an attacker can poison a "
 "shared memory by feeding it crafted inputs, so that later retrievals return the attacker's "
 "content [16]. Ju et al. [42] go further and follow what happens after the first agent is "
 "fooled: manipulated knowledge introduced by a two-stage attack spreads through a community "
 "of language-model agents and stays in circulation as the agents retrieve and reuse it, "
 "rather than remaining with the agent that was first misled. Both results matter here, "
 "because they show a shared memory can carry a corrupted fact from one agent to many. They "
 "also differ from this thesis in one basic way: they assume an attacker who chooses the "
 "false content and intends the harm, while the errors studied here have no attacker behind "
 "them.")
B222 = ("The second strand models how errors move once they are present. Niu et al. [5] adapt "
 "a compartmental epidemic model to a network of language-model agents and derive the "
 "conditions under which errors invade the network. Jamshidi et al. [17] track factual errors "
 "as they pass down a chain of agents and find, on their setup, that the errors shrink rather "
 "than grow. Liu [18] studies how the preferences of a language model acting as an evaluator "
 "spread to other agents. Two further results mark the edges of this space. Gu et al. [43] "
 "show an extreme case. A single crafted input sets off a self-propagating failure that "
 "reaches roughly a million simulated agents and grows exponentially. This is the scale-limit "
 "of what unchecked spread can reach. Shen et al. [44] inject errors into multi-agent systems "
 "wired together in different shapes. They compare each run against a matched run without the "
 "injected error, to measure how much the network's shape changes how far the error travels. "
 "These works share this thesis's view of error spread as a contagion, but they place the "
 "contagion on the agents and the messages between them, not on the facts in a shared memory.")
B223 = ("The third strand tries to contain the spread. Xie et al. [6] inject a single error "
 "into a group of collaborating agents, then add a governance layer that tracks each "
 "message's ancestry. They report that it stops the error from taking over in most runs. "
 "Margalit et al. [19] name the failure modes of shared agent memory, including the loss of "
 "provenance. They build and measure provenance tracking and governed sharing as the fix, but "
 "test it under ordinary operation rather than under a spreading error. Itkin [20] studies the "
 "timing of validation and finds that correcting too late, or too aggressively, can "
 "destabilise the agents' shared belief rather than settle it. The first two build or assume "
 "that provenance-like machinery will contain the spread, which is exactly the claim this "
 "thesis puts to the test.")
CLOSING = ("Section 2.3 compares these three strands with the present work, point by point, "
 "and states what it adds.")
OPT = ("One recent result converges on this thesis's mechanism from a different direction. "
 "Margalit et al. [19], reporting on a production memory service, describe a case where their "
 "duplicate filter runs before their contradiction check. A correcting fact that looks similar "
 "to the one it is meant to fix can then be dropped as a duplicate before anything compares "
 "what it actually says. That is structurally the same problem this thesis finds from the "
 "non-adversarial side: an error that overwrites its own supporting evidence leaves a later "
 "check nothing to catch it with (Chapter 5). That two systems, built for different purposes, "
 "hit the same wall is weak but real evidence that the wall belongs to provenance-governed "
 "shared memory itself, and not to one implementation. The point is tempered by its source, a "
 "single industry report that is not yet peer reviewed and has not been independently "
 "reproduced.")

BIB = [
 (42, 'T. Ju, Y. Wang, X. Ma, P. Cheng, H. Zhao, Y. Wang, L. Liu, J. Xie, Z. Zhang, and '
      'G. Liu, "Flooding Spread of Manipulated Knowledge in LLM-Based Multi-Agent '
      'Communities," arXiv:2407.07791, 2024. Available: ', "https://arxiv.org/abs/2407.07791"),
 (43, 'X. Gu, X. Zheng, T. Pang, C. Du, Q. Liu, Y. Wang, J. Jiang, and M. Lin, "Agent Smith: '
      'A Single Image Can Jailbreak One Million Multimodal LLM Agents Exponentially Fast," in '
      'Proc. 41st Int. Conf. on Machine Learning (ICML), 2024. Available: ',
      "https://arxiv.org/abs/2402.08567"),
 (44, 'X. Shen, Y. Liu, Y. Dai, Y. Wang, R. Miao, Y. Tan, S. Pan, and X. Wang, "Understanding '
      'the Information Propagation Effects of Communication Topologies in LLM-based Multi-Agent '
      'Systems," in Proc. Conf. on Empirical Methods in Natural Language Processing (EMNLP), '
      '2025. Available: ', "https://doi.org/10.18653/v1/2025.emnlp-main.623"),
]

# ---- insert new 2.2 after the heading ----
x = insert_after(heading, INTRO)
x = insert_after(x, "2.2.1 Attacks on agent memory", style="Heading 3")
x = insert_after(x, B221); link_intext(x)
x = insert_after(x, "2.2.2 Models of how errors spread", style="Heading 3")
x = insert_after(x, B222); link_intext(x)
x = insert_after(x, "2.2.3 Attempts to contain the spread", style="Heading 3")
x = insert_after(x, B223); link_intext(x)
x = insert_after(x, CLOSING)

# ---- remove the four old 2.2 paragraphs ----
for old in (old_a, old_b, old_c, old_d):
    old._p.getparent().remove(old._p)

# ---- optional convergent paragraph into 2.3.3 ----
op = insert_after(opt_anchor, OPT); link_intext(op)

# ---- bibliography [42]-[44] ----
bib = bib41; bid = 4200
for n, body, url in BIB:
    bib = insert_after(bib, "[" + str(n) + "] " + body)
    st = OxmlElement("w:bookmarkStart"); st.set(qn("w:id"), str(bid)); st.set(qn("w:name"), "ref"+str(n))
    en = OxmlElement("w:bookmarkEnd"); en.set(qn("w:id"), str(bid))
    if bib._p.pPr is not None: bib._p.pPr.addnext(st)
    else: bib._p.insert(0, st)
    bib._p.append(en); bid += 1
    rid = part.relate_to(url, RT.HYPERLINK, is_external=True)
    base = bib.runs[0]._r
    bib._p.append(hyperlink(run_like(base, url, color="0563C1", underline=True), rid=rid))

d.save(str(DOC))
print("2.2 restructure + 2.3.3 convergent + refs [42]-[44] integrated")
