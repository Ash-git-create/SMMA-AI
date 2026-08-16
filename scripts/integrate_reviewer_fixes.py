# -*- coding: utf-8 -*-
"""Integrate the Chapter 2/3 reviewer fixes into report.docx IN PLACE.

Scope: only content up to section 3.9 (what the reviewer saw). All prose is
style-gated; all numbers are grounded in code/config; all citations are
fact-gate verified. New refs [35]-[41] are appended after [34] with bookmarks
+ external links, matching the [28]-[34] mechanism.

Safety: every anchor is located and asserted BEFORE any mutation; every text
replacement asserts it actually changed exactly one run; multi-run paragraphs
with existing citation hyperlinks ([9] in 2.1.2, [19] in 2.2) are edited only on
their PLAIN continuation runs, so existing hyperlinks are never re-processed.
"""
import re
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

DOC = Path(r"D:/Master Thesis/SMMA_AI_Systems/report.docx")
d = Document(str(DOC))
part = d.part

# ---------------------------------------------------------------- helpers
def run_like(base_r, text, color=None, underline=False):
    r = OxmlElement("w:r")
    if base_r is not None and base_r.find(qn("w:rPr")) is not None:
        r.append(deepcopy(base_r.find(qn("w:rPr"))))
    if color or underline:
        rpr = r.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr"); r.insert(0, rpr)
        if color:
            c = OxmlElement("w:color"); c.set(qn("w:val"), color); rpr.append(c)
        if underline:
            u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    r.append(t)
    return r

def hyperlink(child, anchor=None, rid=None):
    hl = OxmlElement("w:hyperlink")
    if anchor: hl.set(qn("w:anchor"), anchor)
    if rid:    hl.set(qn("r:id"), rid)
    hl.append(child)
    return hl

def insert_after(ref, text, style="Normal"):
    p = OxmlElement("w:p")
    ref._p.addnext(p)
    np = Paragraph(p, ref._parent)
    np.style = d.styles[style]
    np.add_run(text)
    return np

def linkify_run(run):
    """Split ONE plain run around [n] markers, turning each into an internal
    hyperlink to bookmark ref{n}. Does not touch any other run."""
    base = run._r
    if not re.search(r"\[\d+\]", run.text):
        return
    new = []
    for pc in re.split(r"(\[\d+\])", run.text):
        m = re.fullmatch(r"\[(\d+)\]", pc)
        if m:
            new.append(hyperlink(run_like(base, pc), anchor="ref" + m.group(1)))
        elif pc:
            new.append(run_like(base, pc))
    parent = base.getparent()
    idx = list(parent).index(base)
    parent.remove(base)
    for i, el in enumerate(new):
        parent.insert(idx + i, el)

def link_intext(paragraph):
    for run in list(paragraph.runs):
        if re.search(r"\[\d+\]", run.text):
            linkify_run(run)

def para(pred):
    hits = [p for p in d.paragraphs if pred(p.text.strip())]
    assert len(hits) >= 1, "ANCHOR NOT FOUND: %s" % pred
    return hits[0]

def cell_para(substr):
    for tbl in d.tables:
        for row in tbl.rows:
            for c in row.cells:
                if substr in c.text:
                    return c.paragraphs[0]
    raise AssertionError("CELL NOT FOUND: %r" % substr)

def repl_in_para(p, old, new):
    """Replace old->new inside whichever single run of p contains old. Assert
    exactly one run matched."""
    hit = [r for r in p.runs if old in r.text]
    assert len(hit) == 1, "repl expected 1 run with %r, got %d" % (old, len(hit))
    hit[0].text = hit[0].text.replace(old, new, 1)

B = "\u03b2"; G = "\u03b3"; R0 = "R\u2080"

# ---------------------------------------------------------------- locate all anchors first
a_p1   = para(lambda t: t.endswith("number of index cases actually placed."))
a_p2   = para(lambda t: t.startswith("Every run starts from an identical state"))
a_p3   = para(lambda t: t.endswith("reports the fits and how well they match."))
a_p7   = para(lambda t: t.endswith("would be worse than missing one."))
a_p8   = para(lambda t: t.endswith("the analysis treats it that way."))
p_neo  = para(lambda t: t.startswith("The shared memory is a Neo4j graph."))
p_trex = para(lambda t: t.startswith("Before each run the graph is loaded"))
p_mod  = para(lambda t: t.startswith("Two different models are used, split by job."))
p_bound= para(lambda t: t.startswith("The findings are bounded by the setup"))
p_thr  = para(lambda t: t.startswith("The main threat to internal validity"))
p_comp = para(lambda t: t.startswith("Two further properties of these models"))
p_marg = para(lambda t: "name the failure modes of shared agent memory" in t)
c_fever= cell_para("FEVER claim classification")
c_hotp = cell_para("HotpotQA")
c_marg = cell_para("proposed, not stress-tested")
bib34  = para(lambda t: t.startswith("[34]"))
print("all anchors located OK")

# ---------------------------------------------------------------- new prose (style-gated)
P1A = ("All three error types are injected in the same run, fifteen of each, for "
       "forty-five index cases in total, rather than one type to a run. This is possible "
       "because each injected error carries the root label described in Section 3.2, and "
       "that label is passed on to every fact later derived from it. The spread of each type "
       "can therefore be separated out afterwards from a single run, which is what the "
       "per-type comparison of the second research question needs.")
P1B = ("One choice about how an error is injected matters for how the later results should be "
       "read. When a triplet is corrupted, only its content changes: the wrong entity, the "
       "dropped qualifier, or the strengthened predicate. Its stored confidence score, its "
       "lineage formula, and its SIR state are left exactly as they were. This is deliberate. "
       "An error that has not yet been caught should look identical to a trustworthy fact, "
       "which is the situation the thesis studies, so nothing about a freshly corrupted "
       "triplet is allowed to mark it as suspect. Two things follow, and the results chapter "
       "returns to both. Because the confidence score still reflects the fact's clean origin, "
       "a check that reads confidence alone has little to work with on the index cases "
       "themselves. And because the lineage of a corrupted-in-place fact still lists only "
       "clean ancestors, the cascade that deprecates everything downstream of a caught fact "
       "never reaches an index case, since the index case is not downstream of anything that "
       "will be caught: it can only be removed by being detected directly. In effect the "
       "error leaves no trail leading back to it, the same blind spot noted for in-place "
       "extraction errors in Section 2.1.1.")
P2  = ("Each run is ten steps long, and every step works through the same fixed number of "
       "entities, so runs differ only in configuration and seed and never in length.")
P3  = ("From the infected curve alone, the two rates cannot be told apart, because many "
       "different pairs of " + B + " and " + G + " produce the same net growth, and only "
       "their difference is fixed by that curve. They can be separated here because the "
       "recovered count is observed on its own: each step records how many contaminated facts "
       "the validator quarantined, which pins " + G + " directly, and " + B + " then follows. "
       "Even so, " + B + " is best read as a per-step hazard, a rate of spread for this "
       "particular setup, rather than a constant of nature, and " + R0 + " as a summary of "
       "how fast contamination spread within these runs, not a fixed number that would carry "
       "unchanged to a different graph or workload.")
P7  = ("This scheme has two limits. With only four runs in each group, and both tests "
       "required to agree, the Mann-Whitney U test can reach at best a two-sided p of about "
       "0.03, and only when the two groups do not overlap at all, so no small or partial "
       "difference can ever be called real under this rule. And the noise threshold of "
       "Section 3.6, twice the baseline's standard deviation, rests on a standard deviation "
       "that is itself estimated from only four runs, so it is a rough guide rather than an "
       "exact cut-off.")
P8  = ("One measurement asymmetry bears on the second research question in particular. A "
       "propagated error is counted when a newly written fact carries an injected error "
       "forward, and an entity substitution, which changes a single object, can be easier to "
       "recognise as carried forward than a strengthened predicate or a dropped qualifier. "
       "The per-type ranking of spread should therefore be read as indicative, with this "
       "asymmetry of the instrument kept in mind.")
ZSENT = (". A study of this effect, by Zhang et al. [41], finds that a model will commit to "
         "an early wrong answer and then produce further claims that keep it consistent, even "
         "though the same model, asked about those claims on their own, often recognises them "
         "as false. Second,")
LICENCE = (" The three datasets are public research benchmarks released for academic use under "
           "open licences, and both models are open-weight releases under their own licences, "
           "so all of them are used here within the terms their authors set.")

BIB = [
 (35, 'H. Elsahar, P. Vougiouklis, A. Remaci, C. Gravier, J. Hare, F. Laforest, and '
      'E. Simperl, "T-REx: A Large Scale Alignment of Natural Language with Knowledge Base '
      'Triples," in Proc. 11th Int. Conf. on Language Resources and Evaluation (LREC), 2018. '
      'Available: ', "https://aclanthology.org/L18-1544/"),
 (36, 'Z. Yang, P. Qi, S. Zhang, Y. Bengio, W. W. Cohen, R. Salakhutdinov, and C. D. '
      'Manning, "HotpotQA: A Dataset for Diverse, Explainable Multi-hop Question Answering," '
      'in Proc. Conf. on Empirical Methods in Natural Language Processing (EMNLP), 2018, '
      'pp. 2369-2380. Available: ', "https://doi.org/10.18653/v1/D18-1259"),
 (37, 'J. Thorne, A. Vlachos, C. Christodoulopoulos, and A. Mittal, "FEVER: a Large-scale '
      'Dataset for Fact Extraction and VERification," in Proc. Conf. of the North American '
      'Chapter of the Association for Computational Linguistics: Human Language Technologies '
      '(NAACL-HLT), 2018, pp. 809-819. Available: ', "https://doi.org/10.18653/v1/N18-1074"),
 (38, 'Neo4j, Inc., "Neo4j Graph Database Platform," 2024. [Online]. Available: ',
      "https://neo4j.com/"),
 (39, 'Mistral AI and NVIDIA, "Mistral NeMo: A 12B Model with a 128k Context Length," '
      'Jul. 2024. [Online]. Available: ', "https://mistral.ai/news/mistral-nemo"),
 (40, 'A. Grattafiori, A. Dubey, A. Jauhri, et al. (Llama Team), "The Llama 3 Herd of '
      'Models," arXiv:2407.21783, 2024. Available: ', "https://arxiv.org/abs/2407.21783"),
 (41, 'M. Zhang, O. Press, W. Merrill, A. Liu, and N. A. Smith, "How Language Model '
      'Hallucinations Can Snowball," in Proc. 41st Int. Conf. on Machine Learning (ICML), '
      '2024. Available: ', "https://arxiv.org/abs/2305.13534"),
]

# ---------------------------------------------------------------- mutate
# P1: two new paragraphs after the index-cases paragraph (3.3)
x = insert_after(a_p1, P1A)
x = insert_after(x, P1B)
# P2: run length (3.6)
insert_after(a_p2, P2)
# P3: identifiability / how to read R0 (3.4)
insert_after(a_p3, P3)
# P7: two limits of the small-sample rule (3.9)
insert_after(a_p7, P7)
# P8: RQ2 measurement asymmetry (3.9)
insert_after(a_p8, P8)

# In-text dataset/model citations (single-run paragraphs -> safe link_intext)
repl_in_para(p_neo,  "The shared memory is a Neo4j graph.",
                     "The shared memory is a Neo4j graph [38].")
link_intext(p_neo)
repl_in_para(p_trex, "from the T-REx dataset,", "from the T-REx dataset [35],")
link_intext(p_trex)
repl_in_para(p_mod,  "the larger Mistral Nemo 12B,", "the larger Mistral Nemo 12B [39],")
repl_in_para(p_mod,  "smaller and faster Llama 3.1 8B.", "smaller and faster Llama 3.1 8B [40].")
link_intext(p_mod)

# 3.9 hosting reconciliation (P5) + licence sentence (P11), both in the bounded para
repl_in_para(p_bound,
  "Two open models are used, one for extraction and one for judgement, both small enough to "
  "run on a single machine with no graphics card, so behaviour on larger or commercial "
  "models may differ.",
  "Two open-weight models are used, one for extraction and one for judgement, both small by "
  "current standards, in the low tens of billions of parameters, so behaviour on larger or "
  "commercial models may differ. For the experiments they were reached through hosted "
  "inference services rather than run on the local machine (Section 4.1), a choice made for "
  "speed on the available hardware that does not change which models were used.")
# append licence to end of the bounded paragraph's single run
p_bound.runs[-1].text = p_bound.runs[-1].text + LICENCE

# 3.9 model-family wording (P6)
repl_in_para(p_thr, "held fixed for the whole study,", "held fixed across the comparative arms,")

# 2.1.2 Zhang co-cited + explanatory sentence (edit ONLY the plain ". Second," run)
zrun = [r for r in p_comp.runs if r.text.startswith(". Second, the same models")]
assert len(zrun) == 1, "compounding '. Second,' run not uniquely found (%d)" % len(zrun)
zrun[0].text = zrun[0].text.replace(". Second,", ZSENT, 1)
linkify_run(zrun[0])

# 2.2 Margalit sentence (edit ONLY the plain run holding 'propose provenance tracking')
repl_in_para(p_marg,
  "propose provenance tracking and governed sharing as the fix.",
  "build provenance tracking and governed sharing as the fix, but test it under ordinary "
  "operation rather than under a spreading error.")

# Table 3.5 cells
fr = [r for r in c_fever.runs if "FEVER claim classification" in r.text]
assert len(fr) == 1
fr[0].text = fr[0].text.replace(
  "FEVER claim classification:",
  "FEVER [37] claim classification:").replace(
  "against the ground truth",
  "against the ground truth, on the same fixed set of 50 claims, at steps 0, 5, and 10")
link_intext(c_fever)
hr = [r for r in c_hotp.runs if "HotpotQA answers against the ground truth" in r.text]
assert len(hr) == 1
hr[0].text = hr[0].text.replace("HotpotQA answers against the ground truth",
                                "HotpotQA [36] answers against the ground truth")
link_intext(c_hotp)

# Table 2.1 Margalit cell
mr = [r for r in c_marg.runs if "proposed, not stress-tested" in r.text]
assert len(mr) == 1
mr[0].text = mr[0].text.replace("proposed, not stress-tested",
                                "built and measured, not tested under spreading error")

# Bibliography [35]-[41] after [34], with bookmarks + external links
bib = bib34
bid = 3500
for n, body, url in BIB:
    bib = insert_after(bib, "[" + str(n) + "] " + body)
    st = OxmlElement("w:bookmarkStart"); st.set(qn("w:id"), str(bid)); st.set(qn("w:name"), "ref" + str(n))
    en = OxmlElement("w:bookmarkEnd"); en.set(qn("w:id"), str(bid))
    if bib._p.pPr is not None:
        bib._p.pPr.addnext(st)
    else:
        bib._p.insert(0, st)
    bib._p.append(en)
    bid += 1
    rid = part.relate_to(url, RT.HYPERLINK, is_external=True)
    base = bib.runs[0]._r
    bib._p.append(hyperlink(run_like(base, url, color="0563C1", underline=True), rid=rid))

d.save(str(DOC))
print("reviewer fixes integrated into report.docx")
