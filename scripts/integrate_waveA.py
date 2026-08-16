"""Integrate Ch2 Wave A into report.docx in place: insert new 2.1.1 and 2.1.4
paragraphs, link their in-text citations, and add bibliography entries [28]-[31]
(bookmarked + DOI-linked). Preserves existing formatting."""
import re
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

DOC = Path(r"D:/Master Thesis/SMMA_AI_Systems/report.docx")
d = Document(str(DOC))
part = d.part

A = [
    ("How a knowledge graph is built shapes how contamination behaves in it. A knowledge "
     "graph stores each fact as a triple of the form (subject, relation, object), a format "
     "that comes from the Resource Description Framework, a web standard for recording facts "
     "as links between things. Subjects and objects are entities, ideally drawn from a "
     "shared vocabulary so the same thing is always named the same way, and the relation "
     "names the tie between them. Facts enter the graph in one of three ways: entered by "
     "hand, imported from an existing structured source, or extracted from text by a "
     "program. This thesis uses the third route, and it is the one that lets errors in, "
     "because a model reading text can misread it."),
    ("Because extracted graphs contain mistakes, a body of work studies how to find and fix "
     "them, under the name knowledge-graph refinement [28]. Refinement methods look for "
     "facts that are internally inconsistent, that contradict the rest of the graph, or that "
     "are unlikely given the graph's structure. They share an assumption: that an error "
     "leaves a trace the rest of the graph can reveal. One of this thesis's findings is that "
     "a common extraction error, replacing a correct value with a wrong one, leaves no such "
     "trace, because it overwrites the very evidence a refinement check would rely on."),
    ("Sharing a memory between agents is also an old idea. Early artificial-intelligence "
     "systems used a blackboard architecture, in which independent modules read from and "
     "wrote to a common workspace and coordinated only through what they left there, rather "
     "than by messaging each other directly [29]. A shared knowledge graph is a modern "
     "version of that pattern. It brings the same benefit, that any agent can build on any "
     "other's work, and the same risk, that a wrong entry misleads whoever reads it next. "
     "Much of this thesis is a study of that risk in its modern form."),
]
B = [
    ("Database research distinguishes several kinds of provenance, usually summarised as "
     "why, how, and where. Why-provenance names the source records that justify a result; "
     "how-provenance records the way they were combined; and where-provenance points to the "
     "exact place a value was copied from [13]. The lineage this thesis uses is closest to "
     "why- and how-provenance: it records which earlier facts a derived fact rests on and, "
     "through the boolean formula, how they were combined. That is enough to trace a wrong "
     "fact forward to its dependents, which is what the mitigation needs."),
    ("Storing facts with confidences turns a database into a probabilistic database, one "
     "that represents many possible states of the world at once, each with its own "
     "probability [30]. A query over such a database returns, in principle, an answer for "
     "every possible state, weighted by how likely that state is. This is a powerful idea, "
     "and it is also where the cost comes from: the number of possible states grows very "
     "quickly, so exact answers are expensive, which is why practical systems approximate, "
     "as described above."),
    ("A later line of work showed that many of these provenance and confidence calculations "
     "are instances of one algebraic pattern, in which combining facts corresponds to two "
     "operations that behave like addition and multiplication [31]. This is the same pattern "
     "the mitigation uses when it multiplies confidences for an AND and combines them for an "
     "OR. The value of the pattern here is modest but real: it means the simple arithmetic "
     "the mitigation performs rests on an established footing rather than being an arbitrary "
     "choice."),
]
BIB = [
    (28, 'H. Paulheim, "Knowledge Graph Refinement: A Survey of Approaches and Evaluation '
     'Methods," Semantic Web, vol. 8, no. 3, pp. 489-508, 2017. Available: ',
     "https://doi.org/10.3233/SW-160218"),
    (29, 'L. D. Erman, F. Hayes-Roth, V. R. Lesser, and D. R. Reddy, "The Hearsay-II '
     'Speech-Understanding System: Integrating Knowledge to Resolve Uncertainty," ACM '
     'Computing Surveys, vol. 12, no. 2, pp. 213-253, 1980. Available: ',
     "https://doi.org/10.1145/356810.356816"),
    (30, 'D. Suciu, D. Olteanu, C. Re, and C. Koch, "Probabilistic Databases," Synthesis '
     'Lectures on Data Management, Morgan & Claypool, 2011. Available: ',
     "https://doi.org/10.2200/S00362ED1V01Y201105DTM016"),
    (31, 'T. J. Green, G. Karvounarakis, and V. Tannen, "Provenance Semirings," in Proc. ACM '
     'SIGMOD-SIGACT-SIGART Symposium on Principles of Database Systems (PODS), 2007. '
     'Available: ', "https://doi.org/10.1145/1265530.1265535"),
]


def run_like(base_r, text, color=None, underline=False):
    r = OxmlElement("w:r")
    if base_r is not None and base_r.find(qn("w:rPr")) is not None:
        r.append(deepcopy(base_r.find(qn("w:rPr"))))
    if color or underline:
        rpr = r.find(qn("w:rPr")) or OxmlElement("w:rPr")
        if rpr.getparent() is None:
            r.insert(0, rpr)
        if color:
            c = OxmlElement("w:color"); c.set(qn("w:val"), color); rpr.append(c)
        if underline:
            u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    r.append(t)
    return r


def hyperlink(child, anchor=None, rid=None):
    hl = OxmlElement("w:hyperlink")
    if anchor:
        hl.set(qn("w:anchor"), anchor)
    if rid:
        hl.set(qn("r:id"), rid)
    hl.append(child)
    return hl


def insert_after(ref, text):
    p = OxmlElement("w:p")
    ref._p.addnext(p)
    np = Paragraph(p, ref._parent)
    np.style = d.styles["Normal"]
    np.add_run(text)
    return np


def link_intext(paragraph):
    for run in list(paragraph.runs):
        base = run._r
        if not re.search(r"\[\d+\]", run.text):
            continue
        new = []
        for pc in re.split(r"(\[\d+\])", run.text):
            m = re.fullmatch(r"\[(\d+)\]", pc)
            if m:
                new.append(hyperlink(run_like(base, pc), anchor="ref" + m.group(1)))
            elif pc:
                new.append(run_like(base, pc))
        parent = base.getparent()
        idx = list(parent).index(base)
        parent.remove(base)
        for i, el in enumerate(new):
            parent.insert(idx + i, el)


def find(pred):
    return next(p for p in d.paragraphs if pred(p.text.strip()))


# insert 2.1.1 additions
anchor = find(lambda t: t.endswith("a contaminated fact that gets used."))
for txt in A:
    anchor = insert_after(anchor, txt)
    link_intext(anchor)

# insert 2.1.4 additions
anchor = find(lambda t: t.endswith("which is what a live shared memory needs."))
for txt in B:
    anchor = insert_after(anchor, txt)
    link_intext(anchor)

# add bibliography entries [28]-[31] after [27]
bibanchor = find(lambda t: t.startswith("[27]"))
bid = 2000
for n, body, url in BIB:
    bibanchor = insert_after(bibanchor, "[" + str(n) + "] " + body)
    # bookmark ref{n}
    st = OxmlElement("w:bookmarkStart"); st.set(qn("w:id"), str(bid)); st.set(qn("w:name"), "ref" + str(n))
    en = OxmlElement("w:bookmarkEnd"); en.set(qn("w:id"), str(bid))
    if bibanchor._p.pPr is not None:
        bibanchor._p.pPr.addnext(st)
    else:
        bibanchor._p.insert(0, st)
    bibanchor._p.append(en)
    bid += 1
    # external link on the URL: append a hyperlinked run
    rid = part.relate_to(url, RT.HYPERLINK, is_external=True)
    base = bibanchor.runs[0]._r
    bibanchor._p.append(hyperlink(run_like(base, url, color="0563C1", underline=True), rid=rid))

d.save(str(DOC))
print("Wave A integrated into report.docx")
