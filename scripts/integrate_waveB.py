"""Integrate Ch2 Wave B into report.docx in place: append info-cascade bridge
(C1,C2) to end of 2.1.2, insert SIR-assumptions/SEIR (D1,D2) into 2.1.3, link
in-text citations, and add bibliography entries [32]-[34] (bookmarked + DOI-linked).
Preserves existing formatting. Reuses the helpers from the Wave A integrator."""
import re
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

DOC = Path(r"D:/Master Thesis/SMMA_AI_Systems/report.docx")
d = Document(str(DOC))
part = d.part

C = [
    ("Once an error is written into shared memory, the question becomes how far it travels. "
     "When one person passes a claim to others who pass it on again, the result is an "
     "information cascade, and false claims cascade as readily as true ones. A large study "
     "of news spreading on social media found that false stories reached more people and "
     "spread faster than true ones, partly because they were more novel and drew stronger "
     "reactions [32]."),
    ("Studies of spreading behaviour distinguish two ways things spread [33]. In simple "
     "contagion a single exposure is enough to pass something on, as with a cold or a simple "
     "rumour. In complex contagion a person takes something up only after several other "
     "people already have, which is common for beliefs and habits. The contamination this "
     "thesis studies behaves like simple contagion: one retrieval of a wrong fact is enough "
     "for an agent to reuse it, because the agent treats whatever it retrieves as trustworthy "
     "without waiting for a second source to confirm it. That single-exposure behaviour is "
     "what makes the epidemic model of the next section a good fit."),
]
D = [
    ("The basic SIR model makes three simplifying assumptions: a closed population, with no "
     "members entering or leaving during the outbreak; homogeneous mixing, so that every "
     "member is equally likely to meet every other; and a single infected state that "
     "collapses the whole course of an infection into one stage. These assumptions keep the "
     "model simple to fit, and they are reasonable for the short, closed runs this thesis "
     "studies."),
    ("When an infection has a latent period, during which a member is infected but not yet "
     "infectious, a fourth group is added between Susceptible and Infected, giving the SEIR "
     "model, with E for exposed [34]. Other variants let recovered members lose immunity and "
     "turn susceptible again. This thesis keeps the plain SIR form, because a contaminated "
     "fact is either able to spread or has been quarantined, with no meaningful latent stage, "
     "and because the simpler model has fewer parameters to fit from short runs. Section 3.4 "
     "states how the model is fitted here."),
]
BIB = [
    (32, 'S. Vosoughi, D. Roy, and S. Aral, "The Spread of True and False News Online," '
     'Science, vol. 359, no. 6380, pp. 1146-1151, 2018. Available: ',
     "https://doi.org/10.1126/science.aap9559"),
    (33, 'D. Centola, "The Spread of Behavior in an Online Social Network Experiment," '
     'Science, vol. 329, no. 5996, pp. 1194-1197, 2010. Available: ',
     "https://doi.org/10.1126/science.1185231"),
    (34, 'F. Brauer, "Compartmental Models in Epidemiology," in Mathematical Epidemiology, '
     'Lecture Notes in Mathematics, vol. 1945, Springer, 2008, pp. 19-79. Available: ',
     "https://doi.org/10.1007/978-3-540-78911-6_2"),
]


def run_like(base_r, text, color=None, underline=False):
    r = OxmlElement("w:r")
    if base_r is not None and base_r.find(qn("w:rPr")) is not None:
        r.append(deepcopy(base_r.find(qn("w:rPr"))))
    if color or underline:
        rpr = r.find(qn("w:rPr")) or OxmlElement("w:rPr")
        if rpr.getparent() is None:
            r.insert(0, rpr)
        if color:
            c = OxmlElement("w:color"); c.set(qn("w:val"), color); rpr.append(c)
        if underline:
            u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    r.append(t)
    return r


def hyperlink(child, anchor=None, rid=None):
    hl = OxmlElement("w:hyperlink")
    if anchor:
        hl.set(qn("w:anchor"), anchor)
    if rid:
        hl.set(qn("r:id"), rid)
    hl.append(child)
    return hl


def insert_after(ref, text):
    p = OxmlElement("w:p")
    ref._p.addnext(p)
    np = Paragraph(p, ref._parent)
    np.style = d.styles["Normal"]
    np.add_run(text)
    return np


def link_intext(paragraph):
    for run in list(paragraph.runs):
        base = run._r
        if not re.search(r"\[\d+\]", run.text):
            continue
        new = []
        for pc in re.split(r"(\[\d+\])", run.text):
            m = re.fullmatch(r"\[(\d+)\]", pc)
            if m:
                new.append(hyperlink(run_like(base, pc), anchor="ref" + m.group(1)))
            elif pc:
                new.append(run_like(base, pc))
        parent = base.getparent()
        idx = list(parent).index(base)
        parent.remove(base)
        for i, el in enumerate(new):
            parent.insert(idx + i, el)


def find(pred):
    return next(p for p in d.paragraphs if pred(p.text.strip()))


# C1,C2 -> append to end of 2.1.2
anchor = find(lambda t: t.endswith("names the mechanism."))
for txt in C:
    anchor = insert_after(anchor, txt)
    link_intext(anchor)

# D1,D2 -> insert in 2.1.3 after the equations-prose paragraph
anchor = find(lambda t: t.endswith("new cases before it recovers."))
for txt in D:
    anchor = insert_after(anchor, txt)
    link_intext(anchor)

# add bibliography entries [32]-[34] after [31]
bibanchor = find(lambda t: t.startswith("[31]"))
bid = 3200
for n, body, url in BIB:
    bibanchor = insert_after(bibanchor, "[" + str(n) + "] " + body)
    st = OxmlElement("w:bookmarkStart"); st.set(qn("w:id"), str(bid)); st.set(qn("w:name"), "ref" + str(n))
    en = OxmlElement("w:bookmarkEnd"); en.set(qn("w:id"), str(bid))
    if bibanchor._p.pPr is not None:
        bibanchor._p.pPr.addnext(st)
    else:
        bibanchor._p.insert(0, st)
    bibanchor._p.append(en)
    bid += 1
    rid = part.relate_to(url, RT.HYPERLINK, is_external=True)
    base = bibanchor.runs[0]._r
    bibanchor._p.append(hyperlink(run_like(base, url, color="0563C1", underline=True), rid=rid))

d.save(str(DOC))
print("Wave B integrated into report.docx")
