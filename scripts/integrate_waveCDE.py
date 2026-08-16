"""Integrate Waves C/D/E into report.docx in place. No new references: insert
paragraphs and link their in-text [9]/[12]/[16] to existing bookmarks.
- Wave C: one paragraph into 2.1.2 (after the hallucination-causes para).
- Wave D: two paragraphs into 2.1.3 (after the KG-superspreader para).
- Wave E: new section 3.9 (Heading 2 + seven body paragraphs) appended at end of ch3.
Preserves existing formatting."""
import re
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC = Path(r"D:/Master Thesis/SMMA_AI_Systems/report.docx")
d = Document(str(DOC))

C = [
    ("Two further properties of these models matter for what follows. First, their errors "
     "tend to compound: a model shown its own earlier output stays consistent with it, so an "
     "early mistake is carried forward instead of corrected, and a short chain of steps can "
     "drift further from the facts at each one [9]. Second, the same models have no built-in "
     "step that checks a statement against a source of truth before writing it down, so "
     "nothing inside the model stops a plausible-sounding error from being produced [9]. "
     "Extraction, the task at the centre of this thesis, meets both conditions at once, "
     "because it asks the model to turn free text into filled, structured slots at speed."),
]
D = [
    ("It also helps to separate two things an outbreak can be measured by. One is how fast it "
     "grows, which R\u2080 captures. The other is how far it eventually reaches, meaning the "
     "share of the population infected by the time it stops, known as the final size. In the "
     "classic model the two move together, since a higher R\u2080 produces a larger final "
     "size. They come apart when new susceptible members keep arriving, because then even a "
     "slow-spreading process can build up a large cumulative reach over many steps. The "
     "shared memory studied here is closer to that second case, since fresh facts keep "
     "entering it, so this thesis measures both the speed at which contamination spreads and "
     "the total number of facts it eventually touches, because a design can do well on one "
     "and badly on the other."),
    ("Epidemiology also gives a target for control. To stop an infection from growing, enough "
     "of the population must already be immune that each new case reaches, on average, fewer "
     "than one susceptible member. The fraction that must be protected for this to hold is "
     "1 - 1/R\u2080, so a faster-spreading infection, with a higher R\u2080, demands that a "
     "larger share be covered [12]. The same arithmetic applies to a contaminated knowledge "
     "graph. Once validation has already protected part of the graph, the quantity that "
     "matters is the effective reproduction number, meaning the same count measured when not "
     "everyone is still susceptible. Validation has to reach enough of the graph to pull that "
     "number below one, and how much coverage that takes rises with how readily errors "
     "spread. Chapter 5 examines how large that share has to be, and what happens to the "
     "spread when it is not met."),
]
E_HEAD = "3.9 Assumptions, scope, and threats to validity"
E = [
    ("Every experimental design rests on choices that bound what its results can claim. This "
     "section states the main assumptions behind the method, marks what falls inside and "
     "outside its scope, and lists the threats to the validity of its conclusions, together "
     "with the steps taken to limit each one."),
    ("The contamination studied here is non-adversarial. Errors enter because a model "
     "misreads text, not because an attacker plants them, which sets this work apart from "
     "data-poisoning attacks on retrieval systems, where a hostile party crafts inputs to "
     "force a chosen output [16]. Assuming no attacker is the harder case to argue, because "
     "it shows that spread needs no malice, only ordinary model error and a shared memory. It "
     "is also a limit: the results say nothing about how the system behaves under deliberate "
     "attack, and a validator tuned for honest mistakes may well fail against crafted ones."),
    ("The central metric, R\u2080, is borrowed from epidemiology and applied to facts rather "
     "than people. This is an analogy, and it holds only as far as the mapping in Section 3.4 "
     "holds, where a fact is either able to spread, already caught, or not yet reached, with "
     "no state in between. Where the analogy strains, the number should be read as a summary "
     "of spread within these runs, not as a fixed natural constant. Two of the answer-quality "
     "checks, Exact Match and FEVER accuracy, move very little at the contamination levels "
     "reached here (Chapter 5), so their flatness is expected and is not read as evidence "
     "that contamination is harmless. The unsupported sentence ratio of Section 3.7 measures "
     "support, not truth, so a faithful reuse of a contaminated fact counts as supported and "
     "the analysis treats it that way."),
    ("The main threat to internal validity, meaning whether the effect seen is really caused "
     "by what was changed, is the language model's own randomness. A run's seed fixes where "
     "errors are injected, but it does not fix what the models generate, so two runs with the "
     "same seed still differ, and a small gap between single runs is read as noise, not as an "
     "effect. To keep the measuring instrument steady, the judge model used for validation "
     "and auditing is held fixed for the whole study, since changing it partway would "
     "confound a change in the system with a change in the ruler. Because a full run is "
     "expensive on the hardware available, most comparisons rest on the four seeds described "
     "in Section 3.6, so differences in variance are reported as suggestive rather than "
     "settled, and every single-seed figure is labelled as such."),
    ("Small samples are handled conservatively. A comparison between two groups reports both "
     "a Welch t-test, which does not assume the groups have equal spread, and a Mann-Whitney "
     "U test, which does not assume the values follow a normal distribution, and a difference "
     "is called real only when both agree. Comparisons of counts in categories use Fisher's "
     "exact test, which checks whether a split of counts across categories is more lopsided "
     "than chance alone would give. Together with the noise rule of Section 3.6, these "
     "choices trade some ability to detect small effects for caution, which suits a study "
     "where claiming an effect that is not there would be worse than missing one."),
    ("The findings are bounded by the setup that produced them. Two open models are used, one "
     "for extraction and one for judgement, both small enough to run on a single machine with "
     "no graphics card, so behaviour on larger or commercial models may differ. The "
     "pre-loaded facts are drawn from the long tail of a public knowledge base, entities the "
     "models are unlikely to have memorised, which is the fair setting for studying "
     "extraction but may understate how well a model resists an error about a famous entity "
     "it already knows well. The graph is kept in one database engine and queried in one way, "
     "and the three datasets, though varied, do not cover every kind of question. Where a "
     "result is expected to depend on these choices, and where it is expected to survive "
     "them, the analysis says so plainly."),
    ("Finally, the study is built to be repeatable. Every run follows the same three-stage "
     "procedure, loading a fresh graph, extracting into it, then injecting errors and letting "
     "them spread, so that no state leaks from one run into the next. Seeds, configurations, "
     "and per-run outputs are written to files, and the numbers reported in Chapter 5 are "
     "read back from those files, not recalled from notes. This does not remove the model "
     "randomness described above, but it makes every reported figure traceable to an archived "
     "run."),
]


def run_like(base_r, text):
    r = OxmlElement("w:r")
    if base_r is not None and base_r.find(qn("w:rPr")) is not None:
        r.append(deepcopy(base_r.find(qn("w:rPr"))))
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    r.append(t)
    return r


def hyperlink(child, anchor):
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("w:anchor"), anchor)
    hl.append(child)
    return hl


def insert_after(ref, text, style="Normal"):
    p = OxmlElement("w:p")
    ref._p.addnext(p)
    np = Paragraph(p, ref._parent)
    np.style = d.styles[style]
    np.add_run(text)
    return np


def link_intext(paragraph):
    for run in list(paragraph.runs):
        base = run._r
        if not re.search(r"\[\d+\]", run.text):
            continue
        new = []
        for pc in re.split(r"(\[\d+\])", run.text):
            m = re.fullmatch(r"\[(\d+)\]", pc)
            if m:
                new.append(hyperlink(run_like(base, pc), anchor="ref" + m.group(1)))
            elif pc:
                new.append(run_like(base, pc))
        parent = base.getparent()
        idx = list(parent).index(base)
        parent.remove(base)
        for i, el in enumerate(new):
            parent.insert(idx + i, el)


def find(pred):
    return next(p for p in d.paragraphs if pred(p.text.strip()))


# Wave C -> 2.1.2 (after hallucination-causes paragraph)
anchor = find(lambda t: t.endswith("widespread across models and tasks [9]."))
for txt in C:
    anchor = insert_after(anchor, txt)
    link_intext(anchor)

# Wave D -> 2.1.3 (after KG-superspreader paragraph)
anchor = find(lambda t: t.endswith("spread much further than others."))
for txt in D:
    anchor = insert_after(anchor, txt)
    link_intext(anchor)

# Wave E -> new section 3.9 at end of ch3 (after 3.8's last paragraph)
anchor = find(lambda t: t.endswith("not the judge's raw count."))
anchor = insert_after(anchor, E_HEAD, style="Heading 2")
for txt in E:
    anchor = insert_after(anchor, txt)
    link_intext(anchor)

d.save(str(DOC))
print("Waves C/D/E integrated into report.docx")
