"""Make citations clickable in report.docx (in place):
- bookmark each bibliography entry (ref1..refN)
- turn in-text [n] into internal hyperlinks that jump to the matching entry
- turn each bibliography URL into an external hyperlink to the paper
"""
import re
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

DOC = Path(r"D:/Master Thesis/SMMA_AI_Systems/report.docx")
d = Document(str(DOC))
part = d.part


def _run_like(base_r, text, color=None, underline=False):
    r = OxmlElement("w:r")
    if base_r is not None and base_r.find(qn("w:rPr")) is not None:
        r.append(deepcopy(base_r.find(qn("w:rPr"))))
    if color or underline:
        rpr = r.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            r.insert(0, rpr)
        if color:
            c = OxmlElement("w:color")
            c.set(qn("w:val"), color)
            rpr.append(c)
        if underline:
            u = OxmlElement("w:u")
            u.set(qn("w:val"), "single")
            rpr.append(u)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _hyperlink(child_run, anchor=None, rid=None):
    hl = OxmlElement("w:hyperlink")
    if anchor:
        hl.set(qn("w:anchor"), anchor)
    if rid:
        hl.set(qn("r:id"), rid)
    hl.append(child_run)
    return hl


def link_intext(paragraph):
    """Wrap every [n] in the paragraph's runs in an internal hyperlink to ref{n}."""
    for run in list(paragraph.runs):
        base = run._r
        text = run.text
        if not re.search(r"\[\d+\]", text):
            continue
        pieces = re.split(r"(\[\d+\])", text)
        new = []
        for pc in pieces:
            m = re.fullmatch(r"\[(\d+)\]", pc)
            if m:
                new.append(_hyperlink(_run_like(base, pc), anchor="ref" + m.group(1)))
            elif pc:
                new.append(_run_like(base, pc))
        parent = base.getparent()
        idx = list(parent).index(base)
        parent.remove(base)
        for i, el in enumerate(new):
            parent.insert(idx + i, el)


def link_url(paragraph):
    for run in list(paragraph.runs):
        base = run._r
        m = re.search(r"(https?://\S+)", run.text)
        if not m:
            continue
        url = m.group(1)
        before, after = run.text[:m.start()], run.text[m.end():]
        rid = part.relate_to(url, RT.HYPERLINK, is_external=True)
        new = []
        if before:
            new.append(_run_like(base, before))
        new.append(_hyperlink(_run_like(base, url, color="0563C1", underline=True), rid=rid))
        if after:
            new.append(_run_like(base, after))
        parent = base.getparent()
        idx = list(parent).index(base)
        parent.remove(base)
        for i, el in enumerate(new):
            parent.insert(idx + i, el)


# locate bibliography section
paras = d.paragraphs
bib_start = next(i for i, p in enumerate(paras)
                 if p.style.name.startswith("Heading") and p.text.strip() == "Bibliography")
bib_end = len(paras)
for i in range(bib_start + 1, len(paras)):
    if paras[i].style.name.startswith("Heading"):
        bib_end = i
        break

# 1) bookmark + URL-link the bibliography entries
bid = 1000
for i in range(bib_start + 1, bib_end):
    p = paras[i]
    m = re.match(r"^\[(\d+)\]", p.text.strip())
    if not m:
        continue
    name = "ref" + m.group(1)
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    if p._p.pPr is not None:
        p._p.pPr.addnext(start)
    else:
        p._p.insert(0, start)
    p._p.append(end)
    bid += 1
    link_url(p)

# 2) in-text citations: body paragraphs before the bibliography
for i in range(0, bib_start):
    link_intext(paras[i])

# 3) in-text citations inside tables (e.g. Table 2.1)
for t in d.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                link_intext(p)

d.save(str(DOC))
print("citations linked; bibliography bookmarked and URL-linked")
